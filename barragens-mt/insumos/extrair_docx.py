"""Transcreve um .docx para markdown preservando titulos, listas, tabelas e imagens.

Escrito para transcrever insumos oficiais do Vigidesastres sem reescrever o conteudo do
orgao. A prioridade e fidelidade: quando um elemento nao e transcritivel em markdown
(SmartArt, grafico, equacao, imagem sem texto alternativo), o script marca uma lacuna
explicita no ponto correspondente em vez de omitir silenciosamente.

Uso:
    python insumos/extrair_docx.py <origem.docx> <destino.md> [--figuras PASTA]
"""

from __future__ import annotations

import argparse
import datetime as dt
import re
from pathlib import Path
from typing import Any, Iterator

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

# O Word em pt-BR nomeia os estilos de titulo como "Titulo N"; em en-US, "Heading N".
PADRAO_TITULO = re.compile(r"^(?:heading|t[ií]tulo)\s*(\d+)$", re.IGNORECASE)

NS_DIAGRAMA = "http://schemas.openxmlformats.org/drawingml/2006/diagram"
NS_GRAFICO = "http://schemas.openxmlformats.org/drawingml/2006/chart"


class Transcritor:
    def __init__(self, origem: Path, pasta_figuras: Path, prefixo_figuras: str) -> None:
        self.origem = origem
        self.documento = Document(str(origem))
        self.pasta_figuras = pasta_figuras
        self.prefixo_figuras = prefixo_figuras
        self.imagens_salvas: dict[str, str] = {}
        self.lacunas: list[str] = []
        self.contador_imagem = 0

    # ------------------------------------------------------------------ imagens

    def _salvar_imagem(self, rid: str) -> str | None:
        if rid in self.imagens_salvas:
            return self.imagens_salvas[rid]
        try:
            parte = self.documento.part.related_parts[rid]
        except KeyError:
            return None
        self.contador_imagem += 1
        extensao = Path(parte.partname).suffix or ".png"
        nome = f"{self.prefixo_figuras}_fig{self.contador_imagem:02d}{extensao}"
        self.pasta_figuras.mkdir(parents=True, exist_ok=True)
        (self.pasta_figuras / nome).write_bytes(parte.blob)
        self.imagens_salvas[rid] = nome
        return nome

    def _midia_do_paragrafo(self, paragrafo: Paragraph) -> list[str]:
        """Devolve linhas markdown para imagens e marcacoes de lacuna do paragrafo."""
        linhas: list[str] = []
        elemento = paragrafo._p

        for blip in elemento.findall(f".//{qn('a:blip')}"):
            rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
            if not rid:
                continue
            nome = self._salvar_imagem(rid)
            if nome:
                linhas.append(f"![Figura {self.contador_imagem}](figuras/{nome})")
            else:
                linhas.append("> **[LACUNA: imagem vinculada externamente, nao embutida no arquivo]**")
                self.lacunas.append("imagem vinculada externamente")

        for dados in elemento.findall(f".//{qn('a:graphicData')}"):
            uri = dados.get("uri") or ""
            if uri.startswith(NS_DIAGRAMA):
                linhas.append("> **[LACUNA: SmartArt/diagrama nao transcritivel — consultar o .docx original]**")
                self.lacunas.append("SmartArt/diagrama")
            elif uri.startswith(NS_GRAFICO):
                linhas.append("> **[LACUNA: grafico do Word nao transcritivel — consultar o .docx original]**")
                self.lacunas.append("grafico")

        for txbx in elemento.findall(f".//{qn('w:txbxContent')}"):
            texto = " ".join(
                no.text for no in txbx.findall(f".//{qn('w:t')}") if no.text
            )
            if texto.strip():
                linhas.append(f"> **[CAIXA DE TEXTO]** {texto.strip()}")
                self.lacunas.append("caixa de texto (texto recuperado)")

        if elemento.findall(f".//{qn('m:oMath')}") or elemento.findall(f".//{qn('m:oMathPara')}"):
            linhas.append("> **[LACUNA: equacao — ver o .docx original]**")
            self.lacunas.append("equacao")

        return linhas

    # ------------------------------------------------------------- texto inline

    @staticmethod
    def _texto_formatado(paragrafo: Paragraph) -> str:
        partes: list[str] = []
        for execucao in paragrafo.runs:
            texto = execucao.text
            if not texto:
                continue
            # Evita gerar ** ** vazio quando a formatacao cobre apenas espacos.
            if texto.strip():
                esquerda = texto[: len(texto) - len(texto.lstrip())]
                direita = texto[len(texto.rstrip()) :]
                nucleo = texto.strip()
                if execucao.bold:
                    nucleo = f"**{nucleo}**"
                if execucao.italic:
                    nucleo = f"*{nucleo}*"
                texto = f"{esquerda}{nucleo}{direita}"
            partes.append(texto)
        return "".join(partes).strip()

    @staticmethod
    def _nivel_lista(paragrafo: Paragraph) -> int | None:
        pPr = paragrafo._p.pPr
        if pPr is None or pPr.numPr is None:
            return None
        ilvl = pPr.numPr.ilvl
        return int(ilvl.val) if ilvl is not None and ilvl.val is not None else 0

    # ---------------------------------------------------------------- tabelas

    def _texto_celula(self, celula: _Cell) -> str:
        blocos: list[str] = []
        for paragrafo in celula.paragraphs:
            texto = self._texto_formatado(paragrafo)
            if texto:
                blocos.append(texto)
        for tabela in celula.tables:
            blocos.append("[tabela aninhada — ver original]")
            self.lacunas.append("tabela aninhada")
        # Markdown nao aceita quebra de linha real dentro de celula.
        return "<br>".join(blocos).replace("|", "\\|") or " "

    def _tabela_markdown(self, tabela: Table) -> list[str]:
        linhas_texto: list[list[str]] = []
        for linha in tabela.rows:
            try:
                celulas = [self._texto_celula(c) for c in linha.cells]
            except IndexError:
                # Tabelas com celulas mescladas de forma irregular podem falhar aqui.
                self.lacunas.append("linha de tabela com mesclagem irregular")
                continue
            linhas_texto.append(celulas)

        if not linhas_texto:
            return ["> **[LACUNA: tabela vazia ou nao legivel]**"]

        largura = max(len(l) for l in linhas_texto)
        for linha in linhas_texto:
            linha.extend([" "] * (largura - len(linha)))

        saida = ["| " + " | ".join(linhas_texto[0]) + " |",
                 "|" + "---|" * largura]
        for linha in linhas_texto[1:]:
            saida.append("| " + " | ".join(linha) + " |")
        return saida

    # ------------------------------------------------------------------ corpo

    def _blocos(self) -> Iterator[Any]:
        corpo = self.documento.element.body
        for filho in corpo.iterchildren():
            if filho.tag == qn("w:p"):
                yield Paragraph(filho, self.documento)
            elif filho.tag == qn("w:tbl"):
                yield Table(filho, self.documento)

    def transcrever(self) -> str:
        linhas: list[str] = []
        contador_tabela = 0

        for bloco in self._blocos():
            if isinstance(bloco, Table):
                contador_tabela += 1
                linhas.append("")
                linhas.append(f"**Tabela {contador_tabela} do documento original**")
                linhas.append("")
                linhas.extend(self._tabela_markdown(bloco))
                linhas.append("")
                continue

            midia = self._midia_do_paragrafo(bloco)
            texto = self._texto_formatado(bloco)

            if midia:
                linhas.append("")
                linhas.extend(midia)
                linhas.append("")

            if not texto:
                continue

            estilo = (bloco.style.name or "") if bloco.style is not None else ""
            correspondencia = PADRAO_TITULO.match(estilo.strip())
            if correspondencia:
                nivel = min(int(correspondencia.group(1)), 5) + 1
                linhas.append("")
                linhas.append(f"{'#' * nivel} {texto}")
                linhas.append("")
                continue
            if estilo.strip().lower() in {"title", "título", "titulo"}:
                linhas.append("")
                linhas.append(f"# {texto}")
                linhas.append("")
                continue

            nivel_lista = self._nivel_lista(bloco)
            if nivel_lista is not None:
                linhas.append(f"{'  ' * nivel_lista}- {texto}")
                continue

            linhas.append("")
            linhas.append(texto)

        return "\n".join(linhas)

    # --------------------------------------------------------------- cabecalho

    def cabecalho(self, metodo: str) -> str:
        propriedades = self.documento.core_properties
        modificado = self.origem.stat().st_mtime
        secoes = [
            "---",
            f"origem: \"{self.origem}\"",
            f"arquivo: {self.origem.name}",
            f"tamanho_kb: {round(self.origem.stat().st_size / 1024, 1)}",
            f"modificado_em_disco: {dt.datetime.fromtimestamp(modificado).strftime('%Y-%m-%d %H:%M')}",
            f"titulo_documento: {propriedades.title or '(nao informado)'}",
            f"autor: {propriedades.author or '(nao informado)'}",
            f"ultima_modificacao_registrada: {propriedades.modified or '(nao informado)'}",
            f"revisao: {propriedades.revision or '(nao informado)'}",
            f"metodo_extracao: {metodo}",
            f"transcrito_em: {dt.datetime.now().strftime('%Y-%m-%d %H:%M')}",
            "---",
            "",
            "> Transcricao automatica e literal do documento oficial. O conteudo nao foi",
            "> reescrito, resumido nem reordenado. Trechos nao transcritiveis estao marcados",
            "> como **[LACUNA: ...]** no ponto correspondente.",
            "",
        ]
        return "\n".join(secoes)


def main() -> None:
    analisador = argparse.ArgumentParser(description=__doc__)
    analisador.add_argument("origem")
    analisador.add_argument("destino")
    analisador.add_argument("--figuras", default=None)
    analisador.add_argument("--prefixo", default=None)
    argumentos = analisador.parse_args()

    origem = Path(argumentos.origem)
    destino = Path(argumentos.destino)
    pasta_figuras = Path(argumentos.figuras) if argumentos.figuras else destino.parent / "figuras"
    prefixo = argumentos.prefixo or destino.stem[:40]

    transcritor = Transcritor(origem, pasta_figuras, prefixo)
    corpo = transcritor.transcrever()
    cabecalho = transcritor.cabecalho("python-docx 1.2.0 (insumos/extrair_docx.py)")

    destino.parent.mkdir(parents=True, exist_ok=True)
    destino.write_text(cabecalho + corpo.strip() + "\n", encoding="utf-8")

    print(f"gravado {destino}")
    print(f"  {len(corpo.splitlines())} linhas | {transcritor.contador_imagem} imagens")
    if transcritor.lacunas:
        from collections import Counter

        print("  lacunas sinalizadas:")
        for tipo, quantidade in Counter(transcritor.lacunas).most_common():
            print(f"    {quantidade}x {tipo}")


if __name__ == "__main__":
    main()
