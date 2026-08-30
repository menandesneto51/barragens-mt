"""Gera SITREP (relatório de situação) em Markdown / DOCX a partir do recorte do comando."""

from __future__ import annotations

import datetime as dt
import io
from typing import Any

import pandas as pd


def montar_sitrep_md(
    df: pd.DataFrame,
    *,
    municipio: str | None = None,
    proj: dict[str, Any] | None = None,
    tend_clima: str | None = None,
) -> str:
    # Import lazy: evita exigir streamlit ao importar apenas montar_sitrep_cenario_md.
    from st_app.indicadores import indicadores_sanitarios, tendencia_idap_48h

    san = indicadores_sanitarios(df)
    hist = tendencia_idap_48h()
    agora = dt.datetime.now().strftime("%d/%m/%Y %H:%M")
    recorte = municipio or "Estado de Mato Grosso"
    aten = df[df["nivel"].isin(["Amarelo", "Laranja", "Vermelho", "Roxo"])] if not df.empty else df
    top = (
        aten.sort_values("idap_n", ascending=False).head(10)
        if "idap_n" in aten.columns and not aten.empty
        else aten.head(10)
    )

    linhas = [
        f"# SITREP — VIGIBARRAGENS–MT",
        f"**Gerado em:** {agora}",
        f"**Recorte:** {recorte}",
        f"**Barragens no recorte:** {len(df)}",
        "",
        "## 1. Prontidão",
        f"- Em atenção ou pior: **{san['n_atencao']}**",
        f"- População sob pressão sanitária (estimativa): **{san['pop_sob_pressao']:,}**".replace(",", "."),
        f"- US nos municípios sob pressão: **{san['us_sob_risco']}** "
        f"(prioritárias: {san['us_prioritarias']}; método: {san.get('metodo_us', '—')})",
        f"- Municípios sob pressão (sede+jusante): **{san.get('municipios_sob_pressao', san['municipios_jusante'])}**",
        f"- Municípios a jusante distintos: **{san['municipios_jusante']}**",
        f"- Completude média do índice: **{san['completude_media'] if san['completude_media'] is not None else '—'}**",
        "",
        "## 2. Tendência",
    ]
    if hist.get("ok"):
        linhas.append(f"- Índice (últimas rodadas): {hist['msg']}")
    else:
        linhas.append(f"- Índice: {hist.get('msg')}")
    if tend_clima:
        linhas.append(f"- Clima: {tend_clima.replace('**', '')}")
    if proj:
        linhas.append(
            f"- Em atenção+ projetado: {proj.get('amarelo_mais_projetado')} "
            f"({proj.get('delta', 0):+d} vs atual); chuva prevista máx. {proj.get('prevista_max')}"
        )
    linhas += ["", "## 3. Olhar primeiro (até 10)"]
    if top.empty:
        linhas.append("- Nenhuma barragem em atenção no recorte.")
    else:
        for _, r in top.iterrows():
            linhas.append(
                f"- **{r.get('nome','—')}** ({r.get('municipio_sede','—')}) — "
                f"índice {r.get('idap','—')} / {r.get('nivel','—')}"
            )
    linhas += [
        "",
        "## 4. Lacunas operacionais",
        f"- Rejeito/mineração em atenção: {san['rejeito_atencao']}",
        f"- Dano potencial alto sem canal de alerta: {san['dpa_alto_sem_alerta']}",
        f"- Impacto extraterritorial ativo: {san['extraterritorial_ativo']}",
    ]

    # —— ANA / fluvial ——
    try:
        from st_app.ana_fluvial import carregar_estacoes_barragem

        ana = carregar_estacoes_barragem()
        n_acima = 0
        n_cota = 0
        if not ana.empty:
            ids = set()
            if not df.empty and "id_snisb" in df.columns:
                ids = set(df["id_snisb"].astype(str).str.strip())
            sub = ana[ana["id_snisb"].astype(str).str.strip().isin(ids)] if ids else ana
            if "cota_cm" in sub.columns:
                n_cota = int(sub["cota_cm"].notna().sum())
            if "razao_nivel_cota_alerta" in sub.columns:
                raz = pd.to_numeric(sub["razao_nivel_cota_alerta"], errors="coerce")
                n_acima = int((raz >= 1.0).sum())
        linhas += [
            "",
            "## 5. Contexto fluvial (ANA)",
            f"- Estações com cota no recorte: **{n_cota}**",
            f"- Estações ≥ cota de alerta: **{n_acima}**",
            "- Telemetria de rio **não** redefine a mancha proxy.",
        ]
    except Exception:  # noqa: BLE001
        linhas += ["", "## 5. Contexto fluvial (ANA)", "- Indisponível neste ambiente."]

    # —— Contatos ——
    try:
        from st_app.contatos_cobranca import kpis_contatos_criticos, lista_cobranca_contatos

        k = kpis_contatos_criticos(municipio=municipio)
        cob = lista_cobranca_contatos(municipio=municipio)
        linhas += [
            "",
            "## 6. Contatos críticos",
            f"- Papéis críticos com telefone: **{k['n_com_fone']}/{k['n_criticos']}**",
            f"- E-mail pronto para despacho: **{k['email_pronto_despacho']}**",
            f"- Validados ≤90 dias: **{k['n_validados_90d']}**",
            f"- Itens na lista de cobrança: **{k['n_cobranca']}**",
        ]
        if not cob.empty:
            for _, r in cob.head(8).iterrows():
                linhas.append(
                    f"  - {r.get('municipio')} · {r.get('papel_rotulo') or r.get('papel')}: "
                    f"{r.get('motivos')}"
                )
    except Exception:  # noqa: BLE001
        linhas += ["", "## 6. Contatos críticos", "- Cadastro indisponível."]

    # —— PAE ——
    try:
        from st_app.data import ler_csv

        pae = ler_csv("pae_checklist_lacunas.csv")
        n_pae = 0
        if not pae.empty and not df.empty and "id_snisb" in df.columns:
            ids = set(df["id_snisb"].astype(str).str.strip())
            sub = pae[pae["id_snisb"].astype(str).str.strip().isin(ids)]
            col = "n_lacunas_criticas" if "n_lacunas_criticas" in sub.columns else "n_lacuna"
            if col in sub.columns:
                n_pae = int(pd.to_numeric(sub[col], errors="coerce").fillna(0).gt(0).sum())
        linhas += [
            "",
            "## 7. PAE / documentação",
            f"- Barragens do recorte com lacuna PAE: **{n_pae}**",
        ]
    except Exception:  # noqa: BLE001
        linhas += ["", "## 7. PAE / documentação", "- Checklist indisponível."]

    # —— Alertas sem confirmação ——
    try:
        from st_app.ciclo_alerta import resumo_ciclo

        cic = resumo_ciclo()
        linhas += [
            "",
            "## 8. Ciclo de alerta (confirmação)",
            f"- Emitidos: **{cic['n_emitidos']}** · aguardando: **{cic['n_aguardando']}**",
            f"- Confirmados: **{cic['n_confirmados']}** · escalonados: **{cic['n_escalonados']}** "
            f"(máx: {cic['n_escalonado_maximo']})",
        ]
        pend = cic.get("sem_confirmacao")
        if isinstance(pend, pd.DataFrame) and not pend.empty:
            for _, r in pend.head(5).iterrows():
                linhas.append(
                    f"  - {r.get('id_alerta')} · {r.get('nome')} · {r.get('estado')} "
                    f"(limite {r.get('prazo_limite')})"
                )
    except Exception:  # noqa: BLE001
        linhas += ["", "## 8. Ciclo de alerta", "- Sem trilha de ciclo neste ambiente."]

    linhas += [
        "",
        "---",
        "_Proxy geométrico e estimativas rotuladas — não substitui mancha PAE nem ordem de evacuação._",
        "",
    ]
    return "\n".join(linhas)


def montar_sitrep_cenario_md(cenario: dict[str, Any]) -> str:
    """SITREP de um cenário de simulação (mancha ativa + KPIs)."""
    agora = dt.datetime.now().strftime("%d/%m/%Y %H:%M")
    nome = cenario.get("barragem") or "—"
    mun = cenario.get("municipio") or "—"
    geom = cenario.get("geometria") or "—"
    linhas = [
        "# SITREP de cenário — VIGIBARRAGENS–MT",
        f"**Gerado em:** {agora}",
        f"**Barragem:** {nome}",
        f"**Município-sede:** {mun}",
        f"**Geometria ativa:** {geom}",
        "",
        "## 1. Exposição",
        f"- População exposta (setores/proxy): **{cenario.get('pop_exposta', '—')}**",
        f"- Setores na mancha: **{cenario.get('n_setores', '—')}**",
        f"- Captações na mancha: **{cenario.get('n_captacoes', '—')}**",
        f"- Escolas na mancha: **{cenario.get('n_escolas', '—')}**",
        f"- Ativos essenciais (ETA/ETE/energia/abrigos): **{cenario.get('n_ativos', '—')}**",
        f"- MapBiomas urbana sede (ha): **{cenario.get('mapbiomas_ha_urbana', '—')}** "
        f"(drenagem baixa: {cenario.get('mapbiomas_ha_drenagem_baixa', '—')} ha; "
        f"{cenario.get('mapbiomas_pct_drenagem_baixa', '—')}%)",
        "",
        "## 2. Isolamento (C7 proxy)",
        f"- US na mancha / isoladas: **{cenario.get('n_us_atingidas', 0)}** / "
        f"**{cenario.get('n_us_isoladas', 0)}**",
        f"- Vias / pontes: **{cenario.get('n_vias', 0)}** / **{cenario.get('n_pontes', 0)}**",
        f"- Pessoas isoladas (proxy): **{cenario.get('pessoas_isoladas', 0)}**",
        f"- Sedes sem rota / com desvio: **{cenario.get('n_sedes_sem_rota', 0)}** / "
        f"**{cenario.get('n_sedes_com_desvio', 0)}**",
        f"- Desvio médio (km): **{cenario.get('delta_km_medio_desvio', '—')}**",
        f"- Nível C7: **{cenario.get('nivel_c7', '—')}**",
        "",
        "## 3. Clima (dimensão A — ponto da barragem)",
        f"- Chuva 24h / 72h (mm): **{cenario.get('chuva_24h_mm', '—')}** / "
        f"**{cenario.get('chuva_72h_mm', '—')}**",
        f"- Prevista 24–72h (mm): **{cenario.get('chuva_prevista_mm', '—')}**",
        f"- Percentil climatológico: **{cenario.get('percentil_climatologico', '—')}**",
        f"- Fonte telemetria: **{cenario.get('fonte_telemetria', '—')}** "
        f"({cenario.get('aproximacao_espacial', '—')})",
        "",
        "## 4. Capacidade e demanda",
        f"- Pressão estrutural CNES: **{cenario.get('pressao_estrutural', '—')}**",
        f"- Leitos disponíveis (IndicaSUS): **{cenario.get('leitos_disponiveis', '—')}**",
        f"- Demanda internação (2%): **{cenario.get('demanda_internacao', '—')}**",
        f"- Água L/dia (15 L/p): **{cenario.get('demanda_agua', '—')}**",
        f"- IPAPD proxy: **{cenario.get('ipapd', '—')}** ({cenario.get('ipapd_rotulo', '—')}; "
        f"completude {cenario.get('ipapd_completude', '—')})",
        f"- IRS proxy: **{cenario.get('irs', '—')}** ({cenario.get('irs_rotulo', '—')}; "
        f"completude {cenario.get('irs_completude', '—')})",
        "",
        "## 5. PAE / articulação",
        f"- PAE SNISB (PAE-01): **{cenario.get('pae_status', '—')}**",
        f"- Itens checklist lacuna/não: **{cenario.get('pae_lacunas', '—')}**",
        f"- Mancha ZAS oficial: **{cenario.get('pae_zas', '—')}**",
        "",
        "## 6. Contexto fluvial (ANA / SisClima)",
        f"- Estações próximas: **{cenario.get('ana_n_estacoes', '—')}** "
        f"(com cota: {cenario.get('ana_n_com_cota', '—')})",
        f"- Acima da cota de alerta: **{cenario.get('ana_n_acima_alerta', '—')}**",
        f"- A6 com cota medida: **{cenario.get('ana_a6_medido', '—')}**",
        f"- Fonte: `{cenario.get('ana_fonte', '—')}`",
        "",
        "## 7. Contatos / alertabilidade (sede)",
        f"- Cadastro: **{cenario.get('contatos_n', '—')}**",
        f"- Papéis críticos com telefone: **{cenario.get('contatos_criticos', '—')}**",
        f"- Lacunas: **{cenario.get('contatos_faltando', '—')}**",
        "",
        "## 8. Ressalvas",
        "- Proxy geométrico (círculo / trajeto / HAND) — **não** é mancha PAE nem dam break.",
        "- Cota/vazão ANA são contexto operacional e A6 — **não** redimensionam a mancha.",
        "- IPAPD e demanda usam parâmetros a validar; lacunas não são preenchidas com zero.",
        "- MapBiomas é pressão municipal (contexto), não polígono na mancha.",
        "",
    ]
    return "\n".join(linhas)


def montar_sitrep_docx(
    df: pd.DataFrame,
    *,
    municipio: str | None = None,
    proj: dict[str, Any] | None = None,
    tend_clima: str | None = None,
) -> bytes:
    """DOCX de 1 página operacional (python-docx)."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    md = montar_sitrep_md(df, municipio=municipio, proj=proj, tend_clima=tend_clima)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    for linha in md.splitlines():
        if linha.startswith("# "):
            p = doc.add_heading(linha[2:].strip(), level=1)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x1B, 0x32, 0x81)
        elif linha.startswith("## "):
            doc.add_heading(linha[3:].strip(), level=2)
        elif linha.startswith("- "):
            doc.add_paragraph(linha[2:].replace("**", ""), style="List Bullet")
        elif linha.startswith("---"):
            continue
        elif linha.strip():
            doc.add_paragraph(linha.replace("**", ""))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def montar_sitrep_pdf(
    df: pd.DataFrame,
    *,
    municipio: str | None = None,
    proj: dict[str, Any] | None = None,
    tend_clima: str | None = None,
) -> bytes:
    """PDF simples de 1 página (fpdf2)."""
    from fpdf import FPDF

    md = montar_sitrep_md(df, municipio=municipio, proj=proj, tend_clima=tend_clima)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.set_margins(14, 14, 14)
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)
    largura = pdf.epw
    for linha in md.splitlines():
        if linha.startswith("---") or not linha.strip():
            pdf.ln(3)
            continue
        bruto = linha
        if bruto.startswith("# "):
            bruto = bruto[2:]
            pdf.set_font("Helvetica", "B", 14)
            h = 7
        elif bruto.startswith("## "):
            bruto = bruto[3:]
            pdf.set_font("Helvetica", "B", 11)
            h = 6
        else:
            if bruto.startswith("- "):
                bruto = "* " + bruto[2:]
            pdf.set_font("Helvetica", size=10)
            h = 5
        texto = (
            bruto.replace("**", "")
            .replace("—", "-")
            .replace("→", "->")
            .encode("latin-1", "replace")
            .decode("latin-1")
            .strip()
        )
        if not texto:
            continue
        pdf.multi_cell(largura, h, texto)
    out = pdf.output()
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    return bytes(out)
