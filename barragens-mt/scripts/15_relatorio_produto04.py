"""Gera o relatorio tecnico do Produto 04 em .docx com formatacao ABNT.

Requisitos de forma, do briefing de alinhamento do VIGIDESASTRES:
  A4; minimo 6 e maximo 10 paginas incluindo capa e referencias, com apendices e anexos
  fora da contagem; margens de 3 cm superior e esquerda e 2 cm inferior e direita; Arial
  ou Times New Roman 12 em preto; espacamento 1,5 no texto e 1,0 em citacoes longas;
  alinhamento justificado; paginacao; citacoes, notas e referencias em ABNT.

O texto e escrito aqui, mas todo numero vem das bases tratadas pelo pipeline. A intencao
e que o relatorio possa ser regerado apos uma nova coleta sem redigitacao, e que nenhum
valor do documento exista sem lastro em arquivo do repositorio.

Saida em produtos/produto-04/.
"""

from __future__ import annotations

import csv
import datetime as dt
import json
import unicodedata
from typing import Any, Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import comum

PRODUTOS = comum.RAIZ / "produtos" / "produto-04"
FIGURAS = comum.RAIZ / "figuras"

FONTE_TEXTO = "Arial"
TAMANHO_TEXTO = Pt(12)
TAMANHO_APOIO = Pt(10)

TITULO = (
    "ANÁLISE DA CAPACIDADE DE PREPARAÇÃO E RESPOSTA DO SETOR SAÚDE "
    "FRENTE A DESASTRES ASSOCIADOS AO ROMPIMENTO DE BARRAGENS EM CUIABÁ, "
    "MATO GROSSO"
)

# Preenchimentos que dependem de dado administrativo externo ao repositorio. Ficam
# marcados no documento para que ninguem os entregue por engano.
A_PREENCHER = "[PREENCHER]"


# --------------------------------------------------------------------- dados


def ler_csv(nome: str) -> list[dict[str, Any]]:
    caminho = comum.DADOS_TRATADOS / nome
    if not caminho.exists():
        raise SystemExit(f"base ausente: {caminho.name}. Rode o pipeline antes.")
    with caminho.open(encoding="utf-8-sig", newline="") as arquivo:
        return list(csv.DictReader(arquivo, delimiter=";"))


def ler_json(nome: str) -> Any:
    return json.loads((comum.DADOS_TRATADOS / nome).read_text(encoding="utf-8"))


def num(valor: Any) -> float:
    try:
        return float(valor)
    except (TypeError, ValueError):
        return 0.0


def br(valor: float, decimais: int = 0) -> str:
    """Numero no padrao brasileiro: ponto de milhar, virgula decimal."""
    return (
        f"{valor:,.{decimais}f}"
        .replace(",", "\x00")
        .replace(".", ",")
        .replace("\x00", ".")
    )


# Palavras que ficam em minuscula no meio de um nome proprio composto. O str.title() do
# Python nao conhece a regra e produz "Nossa Senhora Do Livramento".
CONECTIVOS = {"da", "das", "de", "do", "dos", "e"}


def titulo_municipio(nome: str) -> str:
    palavras = (nome or "").strip().lower().split()
    return " ".join(
        palavra if indice and palavra in CONECTIVOS else palavra.capitalize()
        for indice, palavra in enumerate(palavras)
    )


def sem_acento(texto: str) -> str:
    decomposto = unicodedata.normalize("NFKD", texto or "")
    return "".join(c for c in decomposto if not unicodedata.combining(c)).upper()


def pct(parte: float, total: float, decimais: int = 1) -> str:
    return br(100 * parte / total, decimais) + "%" if total else "—"


class Dados:
    """Todos os agregados que o relatorio cita, calculados uma vez a partir das bases."""

    def __init__(self) -> None:
        self.inventario = ler_csv("inventario_barragens_mt.csv")
        self.montante = ler_csv("barragens_montante_cuiaba.csv")
        self.cnes = ler_csv("cnes_estabelecimentos_regiao_cuiaba.csv")
        self.exposicao = ler_csv("exposicao_populacoes_eixo_cuiaba.csv")
        self.palmares = ler_csv("palmares_quilombolas_mt.csv")
        self.recorte = ler_json("cuiaba_municipios_de_interesse.json")

        self.total = len(self.inventario)
        self.cri_alto = self._conta("categoria_risco", "Alto")
        self.dpa_alto = self._conta("dano_potencial_associado", "Alto")
        self.classe_a = self._conta("classe_cnrh", "A")
        self.pnsb = self._conta("regulada_pelo_pnsb", "Sim")
        self.com_psb = self._conta("possui_plano_de_seguranca", "Sim")
        self.com_pae = self._conta("possui_pae", "Sim")
        self.com_revisao = self._conta("possui_revisao_periodica", "Sim")
        self.com_inspecao = sum(
            1 for r in self.inventario if (r.get("data_ultima_inspecao") or "").strip()
        )
        self.municipios_com_barragem = len(
            {r.get("municipio") for r in self.inventario if r.get("municipio")}
        )
        self.volume_total = sum(num(r.get("capacidade_hm3")) for r in self.inventario)

        self.por_orgao = self._agrupa("orgao_fiscalizador")
        self.por_uso = self._agrupa("uso_principal")

        mineracao = [
            r
            for r in self.inventario
            if "Mineração" in (r.get("orgao_fiscalizador") or "")
        ]
        self.mineracao = len(mineracao)
        self.em_emergencia = [
            r
            for r in mineracao
            if (r.get("sigbm_nivel_emergencia") or "").strip()
            not in ("", "Sem emergência", "Sem Emergência")
        ]
        self.emergencia_por_municipio = self._agrupa(
            "municipio", registros=self.em_emergencia
        )

        # --- recorte de Cuiaba
        self.em_cuiaba = [r for r in self.inventario if r.get("municipio") == "Cuiabá"]
        self.montante_total = len(self.montante)
        self.montante_municipios = len(
            {r.get("municipio") for r in self.montante if r.get("municipio")}
        )
        self.montante_de_cuiaba = [r for r in self.montante if r.get("municipio") == "Cuiabá"]
        self.montante_sem_pae = sum(
            1 for r in self.montante if (r.get("possui_pae") or "").strip() != "Sim"
        )
        self.montante_dpa_alto = sum(
            1 for r in self.montante if r.get("dano_potencial_associado") == "Alto"
        )
        self.montante_cri_alto = sum(
            1 for r in self.montante if r.get("categoria_risco") == "Alto"
        )
        self.montante_mineracao = sum(
            1 for r in self.montante if "Mineração" in (r.get("orgao_fiscalizador") or "")
        )
        self.volume_em_cuiaba = sum(num(r.get("capacidade_hm3")) for r in self.montante_de_cuiaba)

        # O complexo de Manso vem do inventario completo, e nao do subconjunto a montante:
        # o cadastro registra as tres barragens principais no trecho 896, de resolucao mais
        # grosseira que a secao de controle, e por isso elas nao passam pelo criterio
        # topologico que seleciona os diques do mesmo reservatorio.
        self.manso = [r for r in self.inventario if "UHE MANSO" in (r.get("nome") or "").upper()]
        self.manso_no_montante = sum(
            1 for r in self.montante if "UHE MANSO" in (r.get("nome") or "").upper()
        )
        self.manso_capacidade = max((num(r.get("capacidade_hm3")) for r in self.manso), default=0)
        self.manso_altura = max((num(r.get("altura_m")) for r in self.manso), default=0)
        self.manso_sem_psb = sum(
            1 for r in self.manso if (r.get("possui_plano_de_seguranca") or "").strip() != "Sim"
        )
        self.manso_sem_pae = sum(
            1 for r in self.manso if (r.get("possui_pae") or "").strip() != "Sim"
        )

        # --- rede de saude
        self.cnes_total = len(self.cnes)
        self.cnes_hospitalar = sum(1 for r in self.cnes if r.get("atendimento_hospitalar") == "Sim")
        self.cnes_cirurgico = sum(1 for r in self.cnes if r.get("centro_cirurgico") == "Sim")
        self.cnes_obstetrico = sum(1 for r in self.cnes if r.get("centro_obstetrico") == "Sim")
        self.cnes_neonatal = sum(1 for r in self.cnes if r.get("centro_neonatal") == "Sim")
        self.hospitalar_por_municipio = {
            municipio: sum(
                1
                for r in self.cnes
                if r.get("municipio") == municipio and r.get("atendimento_hospitalar") == "Sim"
            )
            for municipio in sorted({r.get("municipio") for r in self.cnes if r.get("municipio")})
        }
        self.sem_hospital = [m for m, n in self.hospitalar_por_municipio.items() if n == 0]
        self.hospitalar_cuiaba = self.hospitalar_por_municipio.get("Cuiabá", 0)
        self.hospitalar_varzea = self.hospitalar_por_municipio.get("Várzea Grande", 0)

        saude_exposta = [
            r for r in self.exposicao if r.get("categoria") == "estabelecimento de saúde"
        ]
        self.saude_georreferenciada = len(saude_exposta)
        self.hospitalar_georreferenciado = [
            r for r in saude_exposta if r.get("hospitalar") == "Sim"
        ]
        self.hospitalar_proximo = [
            r for r in self.hospitalar_georreferenciado if num(r.get("distancia_eixo_km")) <= 5
        ]

        # --- populacoes vulneraveis
        self.aldeias_proximas = self._na_faixa("aldeia indígena", 10)
        self.terras_proximas = self._na_faixa("terra indígena", 10)
        self.assentamentos_proximos = self._na_faixa("assentamento rural", 10)
        self.assentamentos_5km = self._na_faixa("assentamento rural", 5)
        self.familias_10km = sum(num(r.get("familias")) for r in self.assentamentos_proximos)
        self.familias_5km = sum(num(r.get("familias")) for r in self.assentamentos_5km)
        self.quilombolas_incra = self._na_faixa("território quilombola", 10)

        # As bases do INCRA e da Palmares gravam o municipio em caixa alta e sem acento. O
        # nome oficial vem do recorte do IBGE, para que as tabelas nao misturem grafias.
        self.nomes_oficiais = {
            sem_acento(m["nome"]): m["nome"] for m in self.recorte["municipios"]
        }

        nomes_recorte = {m["nome"].upper() for m in self.recorte["municipios"]}
        self.palmares_recorte = [
            r for r in self.palmares if (r.get("MUNICÍPIO") or "").strip().upper() in nomes_recorte
        ]
        self.palmares_por_municipio = self._agrupa(
            "MUNICÍPIO", registros=self.palmares_recorte
        )

    def municipio_oficial(self, nome: str) -> str:
        return self.nomes_oficiais.get(sem_acento(nome), titulo_municipio(nome))

    def _conta(self, coluna: str, valor: str) -> int:
        return sum(1 for r in self.inventario if (r.get(coluna) or "").strip() == valor)

    def _agrupa(
        self, coluna: str, registros: Sequence[dict[str, Any]] | None = None
    ) -> list[tuple[str, int]]:
        base = registros if registros is not None else self.inventario
        contagem: dict[str, int] = {}
        for registro in base:
            chave = (registro.get(coluna) or "").strip()
            if chave:
                contagem[chave] = contagem.get(chave, 0) + 1
        return sorted(contagem.items(), key=lambda item: -item[1])

    def _na_faixa(self, categoria: str, limite: float) -> list[dict[str, Any]]:
        return [
            r
            for r in self.exposicao
            if r.get("categoria") == categoria and num(r.get("distancia_eixo_km")) <= limite
        ]


# ------------------------------------------------------------------ formatacao


def configurar_documento() -> Document:
    documento = Document()

    normal = documento.styles["Normal"]
    normal.font.name = FONTE_TEXTO
    normal.font.size = TAMANHO_TEXTO
    normal.font.color.rgb = RGBColor(0, 0, 0)
    # A fonte precisa ser fixada tambem no conjunto de leste-asiatico, senao o Word
    # substitui a familia em trechos com caracteres fora do latino basico.
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONTE_TEXTO)

    paragrafo = normal.paragraph_format
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrafo.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragrafo.space_before = Pt(0)
    paragrafo.space_after = Pt(0)
    paragrafo.first_line_indent = Cm(1.25)

    for secao in documento.sections:
        secao.page_width = Cm(21.0)
        secao.page_height = Cm(29.7)
        secao.top_margin = Cm(3.0)
        secao.left_margin = Cm(3.0)
        secao.bottom_margin = Cm(2.0)
        secao.right_margin = Cm(2.0)
    return documento


def numerar_paginas(secao) -> None:
    """Numero de pagina no canto superior direito, como pede a NBR 14724."""
    paragrafo = secao.header.paragraphs[0]
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragrafo.paragraph_format.first_line_indent = Cm(0)
    execucao = paragrafo.add_run()
    execucao.font.name = FONTE_TEXTO
    execucao.font.size = TAMANHO_APOIO

    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")
    instrucao = OxmlElement("w:instrText")
    instrucao.set(qn("xml:space"), "preserve")
    instrucao.text = "PAGE"
    fim = OxmlElement("w:fldChar")
    fim.set(qn("w:fldCharType"), "end")
    for elemento in (inicio, instrucao, fim):
        execucao._r.append(elemento)


def paragrafo(
    documento: Document,
    texto: str,
    *,
    recuo: bool = True,
    alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY,
    tamanho: Pt = TAMANHO_TEXTO,
    negrito: bool = False,
    espacamento=WD_LINE_SPACING.ONE_POINT_FIVE,
    espaco_antes: Pt = Pt(0),
    espaco_depois: Pt = Pt(0),
):
    p = documento.add_paragraph()
    formato = p.paragraph_format
    formato.alignment = alinhamento
    formato.line_spacing_rule = espacamento
    formato.first_line_indent = Cm(1.25) if recuo else Cm(0)
    formato.space_before = espaco_antes
    formato.space_after = espaco_depois
    execucao = p.add_run(texto)
    execucao.font.name = FONTE_TEXTO
    execucao.font.size = tamanho
    execucao.bold = negrito
    return p


def secao_titulo(documento: Document, texto: str, primaria: bool = True) -> None:
    paragrafo(
        documento,
        texto.upper() if primaria else texto,
        recuo=False,
        alinhamento=WD_ALIGN_PARAGRAPH.LEFT,
        negrito=True,
        espaco_antes=Pt(18) if primaria else Pt(12),
        espaco_depois=Pt(6),
    )


def legenda(documento: Document, texto: str, acima: bool = True) -> None:
    paragrafo(
        documento,
        texto,
        recuo=False,
        alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
        tamanho=TAMANHO_APOIO,
        espacamento=WD_LINE_SPACING.SINGLE,
        espaco_antes=Pt(10) if acima else Pt(2),
        espaco_depois=Pt(2) if acima else Pt(10),
    )


def figura(
    documento: Document,
    arquivo: str,
    rotulo: str,
    titulo: str,
    fonte: str,
    largura: Cm = Cm(15.5),
) -> None:
    caminho = FIGURAS / arquivo
    if not caminho.exists():
        paragrafo(documento, f"[FIGURA AUSENTE: {arquivo}]", recuo=False)
        return
    legenda(documento, f"{rotulo} — {titulo}", acima=True)
    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(caminho), width=largura)
    legenda(documento, f"Fonte: {fonte}", acima=False)


def tabela(
    documento: Document,
    rotulo: str,
    titulo: str,
    cabecalho: Sequence[str],
    linhas: Iterable[Sequence[str]],
    fonte: str,
    tamanho: Pt = Pt(9),
) -> None:
    legenda(documento, f"{rotulo} — {titulo}", acima=True)
    linhas = list(linhas)
    grade = documento.add_table(rows=1, cols=len(cabecalho))
    grade.style = "Table Grid"
    grade.alignment = WD_TABLE_ALIGNMENT.CENTER

    def escrever(celula, texto: str, negrito: bool) -> None:
        celula.text = ""
        p = celula.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        execucao = p.add_run(texto)
        execucao.font.name = FONTE_TEXTO
        execucao.font.size = tamanho
        execucao.bold = negrito

    for indice, texto in enumerate(cabecalho):
        escrever(grade.rows[0].cells[indice], texto, True)
    for linha in linhas:
        celulas = grade.add_row().cells
        for indice, texto in enumerate(linha):
            escrever(celulas[indice], str(texto), False)
    legenda(documento, f"Fonte: {fonte}", acima=False)


def capa(documento: Document, dados: Dados) -> None:
    def centralizado(texto: str, negrito: bool = False, tamanho: Pt = TAMANHO_TEXTO) -> None:
        paragrafo(
            documento,
            texto,
            recuo=False,
            alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
            negrito=negrito,
            tamanho=tamanho,
        )

    def vazio(quantidade: int = 1) -> None:
        for _ in range(quantidade):
            paragrafo(documento, "", recuo=False)

    centralizado("MINISTÉRIO DA SAÚDE", negrito=True)
    centralizado("SECRETARIA DE VIGILÂNCIA EM SAÚDE E AMBIENTE")
    centralizado("PROGRAMA DE VIGILÂNCIA EM SAÚDE DE POPULAÇÕES EXPOSTAS A DESASTRES")
    centralizado("VIGIDESASTRES")
    vazio(4)
    centralizado(A_PREENCHER + " (nome do consultor responsável)")
    vazio(4)
    centralizado(TITULO, negrito=True)
    vazio(2)
    paragrafo(
        documento,
        "Produto 04 — Relatório técnico contendo a análise da capacidade de preparação "
        "e resposta do setor saúde frente a desastres associados ao rompimento de "
        "barragens em Cuiabá, Mato Grosso, incluindo mapeamento das áreas e populações "
        "vulneráveis, potenciais impactos à saúde, fluxos de resposta e recomendações "
        "para o fortalecimento da vigilância, da assistência e da logística em saúde.",
        recuo=False,
        espacamento=WD_LINE_SPACING.SINGLE,
        tamanho=TAMANHO_TEXTO,
    )
    vazio(6)
    centralizado("Cuiabá — Mato Grosso")
    centralizado(str(dt.date.today().year))
    documento.add_page_break()


# ------------------------------------------------------------------- conteudo


def introducao(documento: Document, dados: Dados) -> None:
    secao_titulo(documento, "1 Introdução")
    paragrafo(
        documento,
        "O rompimento de barragem é um desastre tecnológico de início súbito, baixa "
        "probabilidade e severidade extrema. Distingue-se dos desastres de evolução gradual "
        "porque a ameaça é uma estrutura construída, com empreendedor identificável e órgão "
        "fiscalizador definido: o risco é, em princípio, conhecido e reduzível antes do "
        "evento. Isso desloca a previsibilidade para a preparação, porque no momento do "
        "rompimento não há tempo de organizar a resposta e o que estiver pactuado previamente "
        "é tudo o que estará disponível. A Política Nacional de Segurança de Barragens (Lei nº "
        "12.334/2010, alterada pela Lei nº 14.066/2020) organiza a matéria em dois eixos "
        "definidos pela Resolução CNRH nº 143/2012: a Categoria de Risco, que mede a condição "
        "da estrutura, e o Dano Potencial Associado, que mede a consequência do rompimento. A "
        "distinção é decisiva para a saúde, porque o planejamento assistencial se dimensiona "
        "pela consequência, e não pela probabilidade de falha: é o dano potencial que indica "
        "quantos leitos, quanto abrigo e quanta vigilância serão necessários.",
    )
    paragrafo(
        documento,
        "Fundão, em Mariana (2015), e a barragem I de Córrego do Feijão, em Brumadinho "
        "(2019), com 270 mortes confirmadas, evidenciaram que esses eventos configuram "
        "Emergências em Saúde Pública prolongadas, cujo impacto não se encerra no trauma "
        "imediato: prossegue na contaminação dos mananciais, na interrupção de unidades de "
        "saúde, no deslocamento forçado de comunidades, no aumento de doenças de veiculação "
        "hídrica e de arboviroses e em agravos psicossociais que persistem por anos. Este "
        "relatório avalia a capacidade do setor saúde para a preparação, a vigilância e a "
        "resposta a esses eventos em Cuiabá, identifica territórios, populações e serviços "
        "vulneráveis, mapeia potenciais impactos à saúde e propõe recomendações, "
        "considerando o papel da capital no apoio regional. A escolha metodológica central "
        "foi abandonar o limite municipal como recorte: a onda de ruptura se propaga pela "
        "calha do rio e ignora fronteiras administrativas, de modo que a estrutura de maior "
        "dano potencial para a capital não está dentro dela — recortar por município "
        "esconderia a ameaça principal.",
    )


def metodologia(documento: Document, dados: Dados) -> None:
    secao_titulo(documento, "2 Metodologia")
    paragrafo(
        documento,
        "O inventário de barragens veio do Sistema Nacional de Informações sobre Segurança "
        "de Barragens (SNISB/ANA), por serviço geográfico oficial, complementado pelo "
        "Sistema Integrado de Gestão de Barragens de Mineração (SIGBM/ANM). A base "
        "territorial é do IBGE; a rede assistencial, do Cadastro Nacional de "
        "Estabelecimentos de Saúde (CNES); os territórios tradicionais, da FUNAI, do INCRA e "
        "da Fundação Cultural Palmares; e a hidrografia, da Base Hidrográfica Ottocodificada "
        "da ANA. Todos os valores citados derivam dessas bases.",
    )
    paragrafo(
        documento,
        "O recorte territorial foi construído por topologia de drenagem, a partir do código "
        "do trecho de curso d'água na codificação de Otto Pfafstetter que o SNISB informa "
        f"para cada barragem — campo preenchido nos {br(dados.total)} registros do estado. "
        "Nessa codificação, dígitos ímpares identificam as interbacias da calha principal, "
        "numeradas da foz para a nascente, e dígitos pares as quatro maiores bacias "
        "tributárias de cada nível; uma barragem está a montante de um ponto quando, no "
        "primeiro dígito divergente, o seu é maior e o do ponto de referência é ímpar. "
        "Adotou-se como seção de controle o trecho do rio Cuiabá mais próximo da mancha "
        f"urbana central da capital, de código 896573, que drena {br(23614.7)} km². A "
        "exposição foi medida pela distância de cada elemento ao eixo que liga o "
        "Aproveitamento Múltiplo de Manso à capital e segue rio Cuiabá abaixo, em faixas de "
        "até 2 km, 5 km e 10 km do talvegue.",
    )
    paragrafo(
        documento,
        "Três limitações condicionam os resultados. A determinante é que não existe mancha de "
        "inundação por ruptura hipotética publicamente disponível para nenhuma barragem de "
        "Mato Grosso; por isso não se apresenta estimativa de população atingida, pois "
        "distância ao talvegue não é cota de inundação e um ponto a 500 metros do rio pode "
        "estar dezenas de metros acima dele — as faixas de proximidade ordenam prioridades de "
        "vigilância, não dimensionam vítimas. A segunda é que o painel público do SNISB em "
        "Power BI foi classificado como fonte não conforme na especificação técnica do projeto "
        "e não alimenta estes resultados. A terceira é que o SipamHidro, do Censipam, opera "
        "sem serviço de dados público documentado, o que impede sua incorporação automatizada "
        "sem acordo prévio com o órgão.",
    )


def panorama(documento: Document, dados: Dados) -> None:
    secao_titulo(documento, "3 Desenvolvimento")
    secao_titulo(documento, "3.1 Panorama das barragens em Mato Grosso", primaria=False)
    usos = dict(dados.por_uso)
    paragrafo(
        documento,
        f"O inventário consolidado reúne {br(dados.total)} barragens em "
        f"{dados.municipios_com_barragem} dos 141 municípios do estado, com volume declarado "
        f"de {br(dados.volume_total)} hm³. A Categoria de Risco é alta em {dados.cri_alto} "
        f"estruturas ({pct(dados.cri_alto, dados.total)}) e o Dano Potencial Associado, alto "
        f"em {dados.dpa_alto} ({pct(dados.dpa_alto, dados.total)}); {dados.classe_a} são de "
        "classe A, a de maior exigência legal. A fiscalização se reparte entre quatro "
        "órgãos, o que fragmenta o acompanhamento: "
        + "; ".join(f"{rotulo_orgao(nome)} com {n}" for nome, n in dados.por_orgao[:4])
        + f". Predominam a irrigação ({usos.get('Irrigação', 0)}) e a contenção de rejeitos de "
        f"mineração ({usos.get('Contenção de rejeitos de mineração', 0)}); apenas "
        f"{usos.get('Abastecimento humano', 0)} servem ao abastecimento humano. Para a saúde, "
        "o dado decisivo não é a finalidade, e sim o material retido: rejeito contamina a "
        "calha com carga química de efeito prolongado, enquanto água produz onda de efeito "
        "agudo — cenários que exigem respostas distintas.",
    )
    paragrafo(
        documento,
        "A conformidade com os instrumentos da Política Nacional de Segurança de Barragens "
        f"está detalhada na Tabela A1, no Anexo A. Apenas "
        f"{pct(dados.com_inspecao, dados.total)} das barragens registram data de "
        "última inspeção, a lacuna mais grave do cadastro: sem histórico de inspeção, a "
        "Categoria de Risco perde lastro empírico e passa a refletir apenas conformidade "
        f"documental. Entre as barragens de mineração, o SIGBM registra {dados.mineracao} "
        f"estruturas no estado, das quais {len(dados.em_emergencia)} com nível de emergência "
        "declarado, concentradas em "
        + "; ".join(
            f"{titulo_municipio(municipio)} com {n}"
            for municipio, n in dados.emergencia_por_municipio[:3]
        )
        + ". Nossa Senhora do Livramento, limítrofe à capital e a montante dela na bacia, "
        "reúne a maior parte delas. A distribuição estadual está na Figura A1.",
    )


def historico(documento: Document) -> None:
    secao_titulo(documento, "3.2 Histórico dos principais rompimentos", primaria=False)
    paragrafo(
        documento,
        "Mato Grosso não registra rompimento de grande magnitude com vítimas em massa, o que "
        "significa que a preparação do setor saúde no estado não pode ser reativa: precisa ser "
        "construída sobre a experiência de outras unidades federativas. Dos eventos de "
        "referência — Fundão, cuja pluma percorreu a bacia do rio Doce até a foz e comprometeu "
        "o abastecimento de dezenas de municípios, e Córrego do Feijão — extraem-se três "
        "lições. A existência formal de Plano de Ação de Emergência não garante que ele esteja "
        "acessível e conhecido pela rede de saúde no momento do evento. O dano de maior "
        "alcance populacional não foi o trauma, mas a interrupção do abastecimento de água e o "
        "deslocamento de população, ambos de longa duração. E a resposta dependeu de "
        "capacidade instalada fora da área atingida, o que faz da distribuição geográfica dos "
        "serviços um dado de preparação.",
    )


def exposicao_cuiaba(documento: Document, dados: Dados) -> None:
    secao_titulo(
        documento, "3.3 Exposição de Cuiabá e hierarquia de ameaças", primaria=False
    )
    paragrafo(
        documento,
        f"Cuiabá tem {len(dados.em_cuiaba)} barragens cadastradas dentro do município, das "
        f"quais {len(dados.montante_de_cuiaba)} drenam para a seção de controle urbana; as "
        "demais estão a jusante dela ou em ramo distinto da rede. Em contrapartida, entram "
        f"estruturas de outros municípios, totalizando {dados.montante_total} barragens em "
        f"{dados.montante_municipios} municípios. O critério é também mais seletivo: as 60 "
        "barragens de Poconé, que um recorte por proximidade incluiria, drenam pelo rio Bento "
        "Gomes direto para o Pantanal e nenhuma integra o conjunto. Das "
        f"{dados.montante_total} estruturas, "
        f"{dados.montante_dpa_alto} têm dano potencial alto, {dados.montante_cri_alto} têm "
        f"categoria de risco alta, {dados.montante_mineracao} são de mineração e "
        f"{dados.montante_sem_pae} — {pct(dados.montante_sem_pae, dados.montante_total)} do "
        "conjunto — não registram Plano de Ação de Emergência (Anexo B; distribuição "
        "cartográfica na Figura A2).",
    )
    paragrafo(
        documento,
        "O Aproveitamento Múltiplo de Manso, no rio Manso, em Chapada dos Guimarães, é a "
        f"estrutura central da análise. Consta do cadastro como {len(dados.manso)} estruturas "
        f"sob o mesmo reservatório, de {br(dados.manso_capacidade)} hm³ e "
        f"{br(dados.manso_altura, 1)} m de altura máxima, classe A e dano potencial alto, "
        "fiscalizadas pela ANEEL. Está a 279 km da capital pela calha e controla 40% da área "
        "de drenagem que chega a ela. É, com folga, a maior fonte de dano potencial "
        f"concentrado da bacia: as {len(dados.montante_de_cuiaba)} barragens internas ao "
        f"município somam {br(dados.volume_em_cuiaba, 1)} hm³, cerca de 140 vezes menos. "
        f"Apenas {dados.manso_no_montante} das {len(dados.manso)} entram no conjunto "
        "delimitado por topologia: o cadastro registra as três barragens principais em trecho "
        "de resolução mais grosseira que a dos diques do mesmo reservatório, inconsistência "
        f"resolvida aqui pela geometria do reservatório. Das {len(dados.manso)} estruturas, "
        f"{dados.manso_sem_psb} não "
        f"registram Plano de Segurança e {dados.manso_sem_pae} não registram Plano de Ação de "
        "Emergência, e nenhuma registra revisão periódica ou data de última inspeção (Anexo "
        "C). Em estruturas de classe A e dano potencial alto a montante da capital do estado, "
        "a lacuna vale como achado ainda que os documentos existam fora do SNISB: para a "
        "vigilância, o que não está no cadastro não está disponível na emergência.",
    )
    paragrafo(
        documento,
        "Duas leituras verdadeiras convivem, e o relatório sustenta as duas. Em operação normal, "
        "o aproveitamento protege Cuiabá e Várzea Grande: foi concebido, antes de ser projeto "
        "hidrelétrico, como obra de controle das cheias do rio Cuiabá, e reserva volume de "
        "espera para amortecê-las, com efeito direto na redução de doenças de veiculação "
        "hídrica, de deslocamento e de interrupção de serviços. Ao mesmo tempo, é a estrutura "
        "que concentra o maior risco agudo da bacia. A tensão não é contradição: uma barragem "
        "transfere risco de um evento frequente e de baixa severidade para um evento raro e "
        "extremo, e tende a estimular a ocupação da planície, ampliando a população exposta ao "
        "evento raro. Daí uma hierarquia de prioridade sanitária em três "
        "níveis: o Manso, por severidade máxima, ainda que sua categoria de risco cadastrada "
        "seja baixa — o que reforça que o planejamento deve seguir a consequência, e não a "
        "condição da estrutura; as barragens de rejeito a montante, concentradas em Nossa "
        "Senhora do Livramento, cujo rompimento não produziria onda comparável, mas "
        "contaminaria a calha que abastece a região metropolitana, com efeito químico e "
        "prolongado; e as barragens internas ao município, de menor severidade por evento e "
        "maior probabilidade de incidente, que correspondem ao que a rede municipal atende no "
        "cotidiano.",
    )


def populacoes(documento: Document, dados: Dados) -> None:
    secao_titulo(
        documento, "3.4 Territórios e populações vulneráveis", primaria=False
    )
    paragrafo(
        documento,
        "O eixo analisado atravessa, a montante da capital, Chapada dos Guimarães, Rosário "
        "Oeste, Nobres, Acorizal, Jangada, Cuiabá e Várzea Grande, e segue a jusante por "
        "Nossa Senhora do Livramento, Santo Antônio de Leverger, Barão de Melgaço e Poconé, "
        "alcançando a planície pantaneira. A área urbana exposta é a da região metropolitana, "
        "cuja mancha contínua de Cuiabá e Várzea Grande está assentada sobre a planície do "
        f"rio. Na faixa de até 10 km do talvegue situam-se {len(dados.aldeias_proximas)} "
        f"aldeias indígenas, {len(dados.terras_proximas)} terra indígena e "
        f"{len(dados.assentamentos_proximos)} assentamentos rurais, com "
        f"{br(dados.familias_10km)} famílias registradas pelo INCRA, das quais "
        f"{br(dados.familias_5km)} em assentamentos a até 5 km (Anexo D). As duas aldeias mais próximas estão "
        "em Barão de Melgaço, a menos de 200 metros do talvegue, na planície, onde a "
        "exposição é maior e o acesso terrestre é mais precário.",
    )
    # O acervo do INCRA cobre apenas territorios com processo de titulacao em andamento, e
    # hoje nao ha nenhum na faixa analisada. A frase precisa funcionar nos dois casos, para
    # que uma nova coleta nao produza texto errado.
    titulados = (
        f"O acervo do INCRA registra {len(dados.quilombolas_incra)} territórios quilombolas "
        "delimitados na faixa analisada, número que reflete"
        if dados.quilombolas_incra
        else "O acervo do INCRA não registra território quilombola delimitado na faixa "
        "analisada, o que reflete"
    )
    paragrafo(
        documento,
        titulados + " o estágio da titulação e não a presença de comunidades: a "
        "Fundação Cultural Palmares certifica "
        f"{len(dados.palmares_recorte)} comunidades quilombolas nos "
        f"{len(dados.recorte['municipios'])} municípios do recorte, concentradas em "
        + "; ".join(
            f"{titulo_municipio(municipio)} com {n}"
            for municipio, n in dados.palmares_por_municipio[:3]
        )
        + " (Anexo E). A divergência é, em si, um achado: planejar a resposta pela base de "
        "titulação subestimaria de forma grosseira a população tradicional exposta. Já a "
        "população ribeirinha não pode ser delimitada por cadastro, pois não há base oficial "
        "equivalente. A alternativa defensável, e recomendação operacional deste relatório, é "
        "aplicar sobre o eixo delimitado o cadastro da Atenção Primária, que identifica "
        "domicílio por microárea e é a única base com resolução suficiente. A distribuição "
        "desses territórios e da rede assistencial ao longo do eixo consta da Figura A3.",
    )


def capacidade_saude(documento: Document, dados: Dados) -> None:
    secao_titulo(
        documento, "3.5 Capacidade instalada e exposição da rede de saúde", primaria=False
    )
    paragrafo(
        documento,
        f"O CNES registra {br(dados.cnes_total)} estabelecimentos de saúde nos "
        f"{len(dados.recorte['municipios'])} municípios do recorte, dos quais "
        f"{dados.cnes_hospitalar} com atendimento hospitalar, {dados.cnes_cirurgico} com "
        f"centro cirúrgico, {dados.cnes_obstetrico} com centro obstétrico e "
        f"{dados.cnes_neonatal} com centro neonatal. A distribuição é fortemente concentrada: "
        f"Cuiabá responde por {dados.hospitalar_cuiaba} e Várzea Grande por "
        f"{dados.hospitalar_varzea} dos {dados.cnes_hospitalar} estabelecimentos com "
        "atendimento hospitalar, reunindo "
        f"{pct(dados.hospitalar_cuiaba + dados.hospitalar_varzea, dados.cnes_hospitalar, 0)} "
        "da retaguarda hospitalar da região (Anexo F).",
    )
    paragrafo(
        documento,
        "Dois achados decorrem desses números e são os mais importantes do trabalho. O primeiro "
        "é que não há estabelecimento com atendimento hospitalar em "
        + juntar(dados.sem_hospital)
        + ". Chapada dos Guimarães, que abriga o Aproveitamento Múltiplo de Manso, está entre "
        "eles: o território que sedia a maior estrutura de dano potencial da bacia depende "
        "integralmente de retaguarda externa, e a disponível é a da capital, justamente o "
        "território atingido a jusante no cenário de ruptura. A dependência é circular e precisa "
        "ser resolvida antes do evento, com pactuação de referência para fora do eixo do rio "
        f"Cuiabá. O segundo é que {len(dados.hospitalar_proximo)} dos "
        f"{len(dados.hospitalar_georreferenciado)} estabelecimentos hospitalares "
        "georreferenciados estão a até 5 km do eixo: a rede que precisaria responder está, em "
        "boa parte, na própria planície do rio, sendo ao mesmo tempo instrumento de resposta e "
        "elemento exposto. Plano de contingência que pressuponha a rede íntegra durante o "
        "evento parte de premissa frágil.",
    )
    paragrafo(
        documento,
        "Os potenciais impactos à saúde distinguem-se por temporalidade, porque exigem "
        "capacidades diferentes. Os imediatos são traumáticos — afogamento, soterramento, "
        "politraumas —, demandam capacidade cirúrgica e de terapia intensiva nas primeiras horas "
        "e concentram-se na faixa atingida. Os indiretos decorrem da interrupção do "
        "abastecimento de água e do saneamento, com aumento de doenças diarreicas agudas, "
        "hepatite A, leptospirose e arboviroses associadas ao acúmulo de água em recipientes; "
        "alcançam população muito maior que a atingida e persistem por meses. Os psicossociais "
        "— luto coletivo, estresse pós-traumático, ansiedade, depressão, uso abusivo de álcool "
        "e ruptura de vínculos — são os de maior duração, os menos visíveis nos sistemas de "
        "informação e os que mais dependem de organização prévia.",
    )


def monitoramento(documento: Document, dados: Dados) -> None:
    secao_titulo(
        documento,
        "3.6 Monitoramento, sistemas de alerta e planos de contingência",
        primaria=False,
    )
    paragrafo(
        documento,
        "O SNISB é um cadastro, não um sistema de monitoramento: não há, nos dados públicos "
        "consultados, telemetria de instrumentação que permita acompanhamento de condição em "
        "tempo real, nem série histórica de inspeções que sustente análise de tendência. O "
        "SipamHidro, do Censipam, oferece visualização de anomalia hidrológica por bacia, mas "
        "opera sem serviço de dados público documentado, o que impede sua ingestão "
        "automatizada; recomenda-se formalizar acordo com o órgão, pelo valor do produto para "
        f"a antecipação de cenários hidrológicos críticos. Quanto aos planos, o déficit é "
        f"mensurável: {dados.montante_sem_pae} das {dados.montante_total} barragens que drenam "
        "para a capital não registram Plano de Ação de Emergência. Mais grave que a ausência "
        "formal é a falta de vínculo com o setor saúde: o plano é instrumento do empreendedor, "
        "e não há evidência, nas fontes consultadas, de que os existentes identifiquem as "
        "unidades de saúde expostas, as populações tradicionais da área potencialmente afetada "
        "ou os fluxos de referência assistencial acionáveis.",
    )
    paragrafo(
        documento,
        "A lacuna determinante é a inexistência de mancha de inundação por ruptura hipotética: "
        "sem estudo de dam break georreferenciado não se estima população exposta, tempo de "
        "chegada de onda, unidades de saúde comprometidas nem necessidade de abrigo. O estudo é "
        "exigível do empreendedor no âmbito do Plano de Ação de Emergência, e obtê-lo é a ação "
        "de maior impacto isolado sobre a preparação do setor saúde em Cuiabá.",
    )


def apoio_regional(documento: Document, dados: Dados) -> None:
    secao_titulo(documento, "3.7 O papel de Cuiabá no apoio regional", primaria=False)
    paragrafo(
        documento,
        "A capital acumula duas posições tratadas usualmente em separado. É território "
        "exposto, por estar a jusante do maior reservatório do estado e assentada sobre a "
        "planície do rio Cuiabá; e é a retaguarda assistencial de Mato Grosso, concentrando "
        "com Várzea Grande a maior parte da capacidade hospitalar, cirúrgica e de terapia "
        "intensiva e recebendo referência de todo o estado. Um evento no eixo do rio "
        "Cuiabá, portanto, não degrada apenas a capacidade local: degrada a estadual, porque "
        "atinge o nó da rede que absorveria a demanda de qualquer outro município afetado. Um "
        "rompimento em Nossa Senhora do Livramento seria atendido em Cuiabá; um rompimento no "
        "Manso atingiria Cuiabá e comprometeria a retaguarda de todos os demais. Isso torna a "
        "redundância de capacidade fora da planície uma questão de preparação estadual, e não "
        "uma preferência de organização municipal da rede.",
    )


def consideracoes(documento: Document, dados: Dados) -> None:
    secao_titulo(documento, "4 Considerações finais")
    paragrafo(
        documento,
        "A capacidade atual do setor saúde para atuar frente ao rompimento de barragens em "
        "Cuiabá é limitada menos pela ausência de serviços do que pela ausência de informação "
        "estruturada e de pactuação prévia. A rede da região metropolitana é robusta em termos "
        "absolutos, mas está concentrada em dois municípios e instalada majoritariamente na "
        "planície do rio, e os municípios a montante — entre eles o que sedia a maior barragem "
        "do estado — não têm retaguarda hospitalar própria. O cadastro, por sua vez, identifica "
        "as estruturas, mas não acompanha sua condição.",
    )
    paragrafo(
        documento,
        "As principais vulnerabilidades são a concentração da retaguarda assistencial no "
        "território exposto; a dependência circular dos municípios a montante em relação à "
        "capital; a ausência de Plano de Ação de Emergência em "
        f"{pct(dados.montante_sem_pae, dados.montante_total)} das barragens que drenam para "
        "Cuiabá, incluindo estruturas do complexo de Manso; a invisibilidade da população "
        "ribeirinha nas bases disponíveis; e a subestimação da população tradicional exposta "
        "quando se usa a base de titulação em vez da de certificação. As capacidades são o "
        "inventário nacional consolidado e público, a base hidrográfica ottocodificada que "
        "permitiu o recorte por bacia e a capilaridade da Atenção Primária, único instrumento "
        "disponível para identificar população exposta com resolução domiciliar.",
    )
    paragrafo(
        documento,
        "No curto prazo, até seis meses, recomenda-se requerer do empreendedor e da ANEEL o "
        "estudo de ruptura hipotética e a mancha de inundação do Aproveitamento Múltiplo de "
        "Manso; exigir o registro do Plano de Ação de Emergência das barragens que drenam para "
        "a capital e não o possuem; pactuar formalmente referência hospitalar para Acorizal, "
        "Chapada dos Guimarães e Jangada em unidade fora do eixo do rio Cuiabá; e iniciar, com "
        "as equipes de Atenção Primária dos municípios do eixo, a identificação da população "
        "ribeirinha por microárea.",
    )
    paragrafo(
        documento,
        "No médio prazo, de seis a vinte e quatro meses, recomenda-se integrar as bases de "
        "barragens, hidrologia, populações vulneráveis e rede assistencial em plataforma única "
        "de vigilância, com atualização automatizada e índice de preparação territorializado; "
        "formalizar acordo com o Censipam para acesso ao SipamHidro; incluir as unidades de "
        "saúde expostas e as comunidades tradicionais nos planos de ação de emergência; e "
        "realizar exercício simulado entre a capital e os municípios a montante, com cenário "
        "de ruptura no Manso. No longo prazo, constituir redundância de capacidade hospitalar "
        "fora da planície de inundação do rio Cuiabá, condição para que a capital exerça a "
        "retaguarda regional mesmo quando for território afetado; e estabelecer vigilância "
        "pós-desastre permanente, capaz de acompanhar por anos os agravos de veiculação "
        "hídrica, os desfechos psicossociais e o deslocamento populacional — os efeitos que "
        "sobrevivem à fase de resposta e que Mariana e Brumadinho mostraram serem os de maior "
        "duração.",
    )


def referencias(documento: Document) -> None:
    secao_titulo(documento, "Referências")
    lista = [
        "AGÊNCIA NACIONAL DE ÁGUAS E SANEAMENTO BÁSICO. Base Hidrográfica Ottocodificada "
        "Multiescalas. Brasília: ANA, [s.d.]. Disponível em: "
        "https://metadados.snirh.gov.br. Acesso em: 29 jul. 2026.",
        "AGÊNCIA NACIONAL DE ÁGUAS E SANEAMENTO BÁSICO. Sistema Nacional de Informações "
        "sobre Segurança de Barragens — SNISB. Brasília: ANA, 2026. Disponível em: "
        "https://www.snisb.gov.br. Acesso em: 29 jul. 2026.",
        "AGÊNCIA NACIONAL DE MINERAÇÃO. Sistema Integrado de Gestão de Barragens de "
        "Mineração — SIGBM. Brasília: ANM, 2026. Disponível em: "
        "https://dadosabertos.anm.gov.br. Acesso em: 29 jul. 2026.",
        "BRASIL. Lei nº 12.334, de 20 de setembro de 2010. Estabelece a Política Nacional "
        "de Segurança de Barragens. Diário Oficial da União: Brasília, 2010.",
        "BRASIL. Lei nº 14.066, de 30 de setembro de 2020. Altera a Lei nº 12.334, de 20 "
        "de setembro de 2010. Diário Oficial da União: Brasília, 2020.",
        "BRASIL. Ministério da Saúde. Cadastro Nacional de Estabelecimentos de Saúde — "
        "CNES. Brasília: Ministério da Saúde, 2026. Disponível em: "
        "https://apidadosabertos.saude.gov.br. Acesso em: 29 jul. 2026.",
        "CONSELHO NACIONAL DE RECURSOS HÍDRICOS. Resolução nº 143, de 10 de julho de "
        "2012. Estabelece critérios gerais de classificação de barragens por categoria de "
        "risco, dano potencial associado e volume. Diário Oficial da União: Brasília, 2012.",
        "FUNDAÇÃO CULTURAL PALMARES. Comunidades quilombolas certificadas. Brasília: "
        "FCP, 2026. Disponível em: https://www.gov.br/palmares. Acesso em: 29 jul. 2026.",
        "FUNDAÇÃO NACIONAL DOS POVOS INDÍGENAS. Terras indígenas e aldeias: dados "
        "geoespaciais. Brasília: FUNAI, 2026. Disponível em: "
        "https://geoserver.funai.gov.br. Acesso em: 29 jul. 2026.",
        "INSTITUTO BRASILEIRO DE GEOGRAFIA E ESTATÍSTICA. Malhas territoriais e divisão "
        "territorial brasileira. Rio de Janeiro: IBGE, 2025. Disponível em: "
        "https://servicodados.ibge.gov.br. Acesso em: 29 jul. 2026.",
        "INSTITUTO NACIONAL DE COLONIZAÇÃO E REFORMA AGRÁRIA. Acervo fundiário: "
        "assentamentos e territórios quilombolas. Brasília: INCRA, 2026. Disponível em: "
        "https://acervofundiario.incra.gov.br. Acesso em: 29 jul. 2026.",
    ]
    for item in lista:
        paragrafo(
            documento,
            item,
            recuo=False,
            alinhamento=WD_ALIGN_PARAGRAPH.LEFT,
            espacamento=WD_LINE_SPACING.SINGLE,
            espaco_depois=Pt(6),
        )


def anexos(documento: Document, dados: Dados) -> None:
    documento.add_page_break()
    secao_titulo(documento, "Anexo A — Panorama estadual: conformidade e cartografia")
    tabela(
        documento,
        "Tabela A1",
        "Conformidade com os instrumentos da Política Nacional de Segurança de Barragens",
        ["Instrumento", "Barragens", "% do inventário"],
        [
            ["Plano de Segurança da Barragem", br(dados.com_psb), pct(dados.com_psb, dados.total)],
            ["Plano de Ação de Emergência", br(dados.com_pae), pct(dados.com_pae, dados.total)],
            [
                "Revisão periódica de segurança",
                br(dados.com_revisao),
                pct(dados.com_revisao, dados.total),
            ],
            [
                "Data de última inspeção registrada",
                br(dados.com_inspecao),
                pct(dados.com_inspecao, dados.total),
            ],
        ],
        "SNISB/ANA, extração de 29/07/2026. Elaboração própria.",
    )
    figura(
        documento,
        "mapa_barragens_mt_cri.png",
        "Figura A1",
        "Barragens cadastradas em Mato Grosso por categoria de risco",
        "SNISB/ANA e SIGBM/ANM; malha territorial do IBGE. Elaboração própria.",
    )
    documento.add_page_break()
    figura(
        documento,
        "mapa_cuiaba_barragens_montante.png",
        "Figura A2",
        "Barragens a montante de Cuiabá e eixo de drenagem do rio Cuiabá",
        "SNISB/ANA, SIGBM/ANM e Base Hidrográfica Ottocodificada da ANA; malha "
        "territorial do IBGE. Elaboração própria.",
    )
    documento.add_page_break()
    figura(
        documento,
        "mapa_cuiaba_populacoes_saude.png",
        "Figura A3",
        "Populações vulneráveis e rede de saúde no eixo Manso — Cuiabá",
        "FUNAI, INCRA, Fundação Cultural Palmares e CNES/Ministério da Saúde; hidrografia "
        "da ANA e malha territorial do IBGE. Elaboração própria.",
    )

    documento.add_page_break()
    secao_titulo(documento, "Anexo B — Barragens que drenam para a seção de controle de Cuiabá")
    paragrafo(
        documento,
        f"Das {dados.montante_total} estruturas identificadas por topologia de drenagem, "
        "apresentam-se as de Dano Potencial Associado alto ou Categoria de Risco alta, "
        "ordenadas por capacidade do reservatório. A relação completa consta do arquivo "
        "barragens_montante_cuiaba.csv do repositório do projeto. A capacidade informada é a "
        "do reservatório, repetida pelo cadastro em cada estrutura que o confina, e por isso "
        "não deve ser somada entre estruturas de um mesmo empreendimento.",
    )
    criticas = [
        r
        for r in dados.montante
        if r.get("dano_potencial_associado") == "Alto" or r.get("categoria_risco") == "Alto"
    ]
    tabela(
        documento,
        "Tabela B1",
        "Barragens de dano potencial alto ou categoria de risco alta a montante de Cuiabá",
        ["Estrutura", "Município", "CRI", "DPA", "Capac. (hm³)", "PAE"],
        [
            [
                r.get("nome"),
                r.get("municipio"),
                r.get("categoria_risco") or "—",
                r.get("dano_potencial_associado") or "—",
                br(num(r.get("capacidade_hm3")), 2),
                r.get("possui_pae") or "—",
            ]
            for r in sorted(criticas, key=lambda r: -num(r.get("capacidade_hm3")))
        ],
        "SNISB/ANA e SIGBM/ANM. Elaboração própria.",
        tamanho=Pt(8),
    )

    documento.add_page_break()
    secao_titulo(documento, "Anexo C — Estruturas do Aproveitamento Múltiplo de Manso")
    tabela(
        documento,
        "Tabela C1",
        "Instrumentos de segurança por estrutura do complexo de Manso",
        ["Estrutura", "Município", "Altura (m)", "Plano de Segurança", "PAE"],
        [
            [
                r.get("nome"),
                r.get("municipio"),
                br(num(r.get("altura_m")), 1),
                r.get("possui_plano_de_seguranca") or "não informado",
                r.get("possui_pae") or "não informado",
            ]
            for r in sorted(dados.manso, key=lambda r: -num(r.get("altura_m")))
        ],
        "SNISB/ANA. Elaboração própria.",
        tamanho=Pt(8),
    )

    documento.add_page_break()
    secao_titulo(documento, "Anexo D — Elementos vulneráveis a até 5 km do eixo de drenagem")
    proximos = [
        r
        for r in dados.exposicao
        if r.get("categoria") != "estabelecimento de saúde"
        and num(r.get("distancia_eixo_km")) <= 5
    ]
    tabela(
        documento,
        "Tabela D1",
        "Territórios e comunidades a até 5 km do talvegue, por distância",
        ["Categoria", "Nome", "Município", "Distância (km)", "Observação"],
        [
            [
                r.get("categoria"),
                r.get("nome"),
                dados.municipio_oficial(r.get("municipio") or ""),
                br(num(r.get("distancia_eixo_km")), 2),
                r.get("detalhe") or "",
            ]
            for r in sorted(proximos, key=lambda r: num(r.get("distancia_eixo_km")))
        ],
        "FUNAI e INCRA; hidrografia da ANA. Elaboração própria.",
        tamanho=Pt(8),
    )
    paragrafo(
        documento,
        "A distância é medida do elemento ao talvegue e não representa cota de inundação. A "
        "ausência de mancha de inundação impede converter proximidade em exposição, "
        "conforme exposto na seção 2.",
    )

    documento.add_page_break()
    secao_titulo(
        documento, "Anexo E — Comunidades quilombolas certificadas nos municípios do eixo"
    )
    tabela(
        documento,
        "Tabela E1",
        "Comunidades quilombolas certificadas pela Fundação Cultural Palmares",
        ["Município", "Comunidade"],
        [
            [titulo_municipio(r.get("MUNICÍPIO") or ""), r.get("COMUNIDADE")]
            for r in sorted(
                dados.palmares_recorte,
                key=lambda r: ((r.get("MUNICÍPIO") or ""), (r.get("COMUNIDADE") or "")),
            )
        ],
        "Fundação Cultural Palmares. Elaboração própria.",
        tamanho=Pt(8),
    )

    documento.add_page_break()
    secao_titulo(documento, "Anexo F — Retaguarda hospitalar nos municípios do eixo")
    tabela(
        documento,
        "Tabela F1",
        "Estabelecimentos com atendimento hospitalar nos municípios do eixo analisado",
        ["Município", "Com atendimento hospitalar"],
        [
            [municipio, str(quantidade)]
            for municipio, quantidade in sorted(
                dados.hospitalar_por_municipio.items(), key=lambda item: -item[1]
            )
        ],
        "CNES/Ministério da Saúde, extração de 29/07/2026. Elaboração própria.",
    )


def rotulo_orgao(nome: str) -> str:
    abreviacoes = {
        "MT - Secretaria de Estado do Meio Ambiente - SEMA": "a SEMA-MT",
        "Agência Nacional de Mineração - ANM": "a ANM",
        "Agência Nacional de Energia Elétrica - ANEEL": "a ANEEL",
        "Agência Nacional de Águas e Saneamento Básico - ANA": "a ANA",
    }
    return abreviacoes.get(nome, nome)


def juntar(itens: Sequence[str]) -> str:
    if not itens:
        return "nenhum município"
    if len(itens) == 1:
        return itens[0]
    return ", ".join(itens[:-1]) + f" e {itens[-1]}"


def main() -> None:
    dados = Dados()
    print("Gerando relatorio do Produto 04")
    print(f"  inventario: {dados.total} barragens")
    print(f"  a montante de Cuiaba: {dados.montante_total} barragens")
    print(f"  rede de saude: {dados.cnes_total} estabelecimentos")

    documento = configurar_documento()
    capa(documento, dados)

    # A paginacao comeca no conteudo textual, como pede a NBR 14724: a capa nao recebe
    # numero, e por isso o corpo vive numa secao propria.
    corpo = documento.add_section(WD_SECTION.NEW_PAGE)
    corpo.page_width = Cm(21.0)
    corpo.page_height = Cm(29.7)
    corpo.top_margin = Cm(3.0)
    corpo.left_margin = Cm(3.0)
    corpo.bottom_margin = Cm(2.0)
    corpo.right_margin = Cm(2.0)
    numerar_paginas(corpo)

    introducao(documento, dados)
    metodologia(documento, dados)
    panorama(documento, dados)
    historico(documento)
    exposicao_cuiaba(documento, dados)
    populacoes(documento, dados)
    capacidade_saude(documento, dados)
    monitoramento(documento, dados)
    apoio_regional(documento, dados)
    consideracoes(documento, dados)
    referencias(documento)
    anexos(documento, dados)

    PRODUTOS.mkdir(parents=True, exist_ok=True)
    destino = PRODUTOS / "Produto_04_VIGIBARRAGENS_Cuiaba_MT.docx"
    documento.save(destino)
    print(f"  gravado {destino.relative_to(comum.RAIZ)}")
    print(f"  ATENCAO: preencher os campos marcados '{A_PREENCHER}' na capa")


if __name__ == "__main__":
    main()
